# import os
# import json
# import re
# from flask import Flask, render_template, request, jsonify, send_from_directory
# import fitz  # PyMuPDF
# import docx
# import faiss
# import numpy as np
# from sentence_transformers import SentenceTransformer
# import ollama  # CHANGED: Imported ollama
# from werkzeug.utils import secure_filename

# # --- Global App Settings & Model Paths ---
# # CHANGED: Replaced GGUF path with Ollama model name
# ollama_model_name = "llama3.1:8b-instruct-q8_0" 

# embedding_model_path = "/Users/abhishekprasad/Documents/RAG/Embedding_Models/E5-Large"

# # --- Paths for the ingested style guide ---
# FAISS_INDEX_PATH = "data/style_guide.faiss"
# CHUNKS_JSON_PATH = "data/style_guide_chunks.json"

# # --- Global Variables ---
# style_guide_faiss_index = None
# style_guide_text_chunks = []
# embedding_model = None
# llm = None  # This will now hold our Ollama wrapper function
# reviser_llm = None
# AGENT_CATEGORIES = ['Capitalization', 'Punctuation', 'Text Formatting']

# # --- Setup Flask and Directories ---
# app = Flask(__name__)
# UPLOADS_DIR = "uploads"
# GENERATED_DIR = "generated" # NEW: Directory for commented files
# os.makedirs(UPLOADS_DIR, exist_ok=True)
# os.makedirs(GENERATED_DIR, exist_ok=True)
# ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

# # --- Load Models and Style Guide Index on Startup ---
# def load_resources():
#     global embedding_model, reviser_llm, llm, style_guide_faiss_index, style_guide_text_chunks
#     try:
#         print("Loading embedding model...")
#         embedding_model = SentenceTransformer(embedding_model_path)
        
#         # --- CHANGED: Replaced Llama_cpp init with Ollama wrapper ---
#         print("Creating Ollama client function...")

#         def ollama_client_function(prompt, max_tokens, temperature, stop):
#             """
#             This function mimics the llama_cpp llm() call signature
#             and return structure to avoid changing other parts of the code.
#             """
#             try:
#                 # Note: llama_cpp 'max_tokens' is 'num_predict' in ollama options
#                 response = ollama.generate(
#                     model=ollama_model_name,
#                     prompt=prompt,
#                     stream=False,
#                     options={
#                         'num_ctx': 35000,
#                         'temperature': temperature,
#                         'stop': stop,
#                         'num_predict': max_tokens,
#                         'seed': 42  # Maintain seed for consistency
#                     }
#                 )
#                 # Mimic the llama_cpp output structure {'choices': [{'text': ...}]}
#                 return {'choices': [{'text': response['response']}]}
#             except Exception as e:
#                 print(f"Error calling Ollama: {e}")
#                 # Return a structure that won't crash the next line
#                 return {'choices': [{'text': '[]'}]} # Return empty JSON list on error

#         llm = ollama_client_function  # Assign the function to the global 'llm' variable
        
#         try:
#             ollama.list()
#             print("Ollama service connected successfully.")
#         except Exception as e:
#             print(f"WARNING: Could not connect to Ollama service. Please ensure Ollama is running.")
#             print(f"Error: {e}")
#         # --- End of Change ---

#         print("Loading Style Guide FAISS index...")
#         style_guide_faiss_index = faiss.read_index(FAISS_INDEX_PATH)
#         print("Loading Style Guide text chunks...")
#         with open(CHUNKS_JSON_PATH, "r", encoding="utf-8") as f:
#             style_guide_text_chunks = json.load(f)
#         print("All resources loaded successfully!")
#     except Exception as e:
#         print(f"CRITICAL ERROR during resource loading: {e}")

# load_resources()

# # --- Utility Functions ---
# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf_fitz(file_path):
#     text = ""
#     try:
#         doc = fitz.open(file_path)
#         for page in doc:
#             text += page.get_text("text") + "\n"
#         doc.close()
#         # Clean up common PDF extraction artifacts
#         text = re.sub(r'\s*\n\s*', '\n', text).strip()
#         return text
#     except Exception as e:
#         print(f"Error extracting from PDF with PyMuPDF: {e}")
#         return ""

# def extract_text_from_docx(file_path):
#     text = []
#     try:
#         doc = docx.Document(file_path)
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 text.append(para.text.strip())
#         return "\n\n".join(text)
#     except Exception as e:
#         print(f"Error extracting from DOCX: {e}")
#         return ""

# def parse_llm_json_output(llm_response_text):
#     match = re.search(r'```json\s*([\s\S]*?)\s*```', llm_response_text)
#     json_str = match.group(1) if match else llm_response_text
#     last_bracket_pos = json_str.rfind(']')
#     if last_bracket_pos != -1:
#         json_str = json_str[:last_bracket_pos + 1]
#     else: return []
#     try:
#         parsed_json = json.loads(json_str)
#         return [parsed_json] if isinstance(parsed_json, dict) else parsed_json
#     except json.JSONDecodeError: return []

# def run_reviser_agent(correction, relevant_rules_text):
#     """
#     Uses a smaller LLM to act as a 'Reviser'. It refines, approves, or rejects a proposed correction.
#     Returns a valid correction dictionary if approved/revised, or None if rejected.
#     """
#     original = correction.get("original_text")
#     corrected = correction.get("corrected_text")
#     reason = correction.get("consolidated_reason")

#     system_prompt = """You are a meticulous senior editor. A junior AI agent has proposed a correction based on the provided 'Style Guide Rules'.
# Your task is to review this proposal and act as a final quality gate.

# 1.  **Analyze:** Carefully compare the 'original_text' to the 'corrected_text' and evaluate the 'reason' against the 'Style Guide Rules'.
# 2.  **Decision:**
#     - If the correction is perfect and well-justified by the rules, output the original JSON object.
#     - If the correction is mostly correct but the reason is weak or the text could be slightly better, **refine** the `corrected_text` or `consolidated_reason` and output the updated JSON object.
#     - If the correction is incorrect, unnecessary, or not supported by the provided rules, you MUST reject it by responding with an empty JSON object `{}`.

# **Your entire output must be a single, valid JSON object.** Do not add any commentary or text outside the JSON structure.
# """
    
#     user_question = f"""Style Guide Rules:
# ---
# {relevant_rules_text}
# ---

# Proposed Correction (JSON):
# ```json
# {{
#   "original_text": "{original}",
#   "corrected_text": "{corrected}",
#   "consolidated_reason": "{reason}",
#   "issues_found": {correction.get("issues_found", [])}
# }}
# Review the proposal and respond with the final, approved/revised JSON object, or an empty object {{}} to reject it.
# """

# def run_triage_agent(text_batch):
#     categories_str = ", ".join(AGENT_CATEGORIES)
#     system_prompt = f"Analyze the following text. Identify which of these style categories are relevant for an editorial review: {categories_str}. Respond ONLY with a comma-separated list of the relevant categories. If none are relevant, respond with 'None'."
#     #----llama prompt----
#     prompt = f"<|start_header_id|>system<|end_header_id|>\n\n{system_prompt}<|eot_id|><|start_header_id|>user<|end_header_id|>\n\nText: \"{text_batch}\"\n\nRelevant Categories:<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n"
    
#     try:
#         # NO CHANGE NEEDED: llm() call is handled by our wrapper
#         result = llm(prompt, max_tokens=50, temperature=0.0, stop=["<|eot_id|>"])
#         response_text = result['choices'][0]['text'].strip()
#         if "None" in response_text: return []
#         return [cat.strip() for cat in response_text.split(',') if cat.strip() in AGENT_CATEGORIES]
#     except Exception: return []

# def retrieve_relevant_rules(query_text, top_k=10):
#     query_embedding = embedding_model.encode([query_text]).astype(np.float32)
#     _, indices = style_guide_faiss_index.search(query_embedding, min(top_k, len(style_guide_text_chunks)))
#     return [style_guide_text_chunks[i] for i in indices[0]]

# def generate_review_for_batch(text_batch):
#     print(f"\n--- Analyzing Batch ---\n{text_batch[:300]}...")

#     print("Retrieval: Searching for relevant rules based on content...")
#     relevant_rules = retrieve_relevant_rules(text_batch, top_k=10) # Increased top_k slightly for broader context

#     if not relevant_rules:
#         print("Retrieval: No relevant style guide rules found for this batch.")
#         return []

#     MAX_CONTEXT_CHARS = 4000  # Max characters for the retrieved rules
#     context_parts = []
#     current_context_chars = 0

#     unique_rules = []
#     seen_rules = set()
#     for rule_dict in relevant_rules:
#         rule_text = rule_dict.get('full_text', '')
#         if rule_text and rule_text not in seen_rules:
#             if current_context_chars + len(rule_text) <= MAX_CONTEXT_CHARS:
#                 unique_rules.append(rule_text)
#                 seen_rules.add(rule_text)
#                 current_context_chars += len(rule_text)
#             else:
#                 break # Stop if we exceed the context character limit
            
#     context = "\n\n---\n\n".join(unique_rules)
#     print(f"Retrieval: Built context with {len(unique_rules)} rules ({current_context_chars} chars).")

#     system_prompt = """You are an expert editor for the Microsoft Writing Style Guide. Your sole task is to analyze the user's text based ONLY on the provided 'Style Guide Rules'. Do not use any outside knowledge.

# Your goal is to identify clear violations and provide corrections.

# **CRITICAL RULES TO FOLLOW:**
# 1.  **High Confidence Only:** Propose a correction only when you have high confidence that a rule is explicitly being broken. If a sentence is stylistically awkward but doesn't violate a provided rule, ignore it.
# 2.  **No Empty Changes:** Only include an object in the JSON list if the `corrected_text` is different from the `original_text`.
# 3.  **Strictly No Output for No Errors:** If a sentence or segment has NO clear violations of the provided rules, DO NOT create a JSON object for it. It should be completely omitted from your response.
# 4.  **Valid JSON Output:** Your entire output must be a single, valid JSON list of objects.

# Each object in the list must have exactly four keys:
# 1. "original_text": The full original sentence or text segment.
# 2. "corrected_text": The fully corrected version of the sentence, with ALL issues fixed according to the rules.
# 3. "issues_found": A list of short strings naming the issue types found (e.g., ["Capitalization", "Punctuation"]).
# 4. "consolidated_reason": A single, brief explanation summarizing all the changes made and citing the relevant rule(s).

# If you find zero violations of the provided rules in the entire text batch, you MUST respond with an empty list `[]`.
# """
    
#     user_question = f"Style Guide Rules:\n---\n{context}\n---\n\nUser Text to Review:\n---\n{text_batch}\n---\n\nAnalyze the text and provide the consolidated corrections in the specified JSON format."
#     prompt = f"<|start_header_id|>system<|end_header_id|>\n\n{system_prompt}<|eot_id|><|start_header_id|>user<|end_header_id|>\n\n{user_question}<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n"
       
#     try:
#         # NO CHANGE NEEDED: llm() call is handled by our wrapper
#         result = llm(prompt, max_tokens=1000, temperature=0.0, stop=["<|eot_id|>"])
#         response_text = result['choices'][0]['text'].strip()
#         corrections = parse_llm_json_output(response_text)
#         print(f"Analysis: Found {len(corrections)} consolidated corrections in batch.")
#         return corrections, context
#     except Exception as e:
#         print(f"Error in Analysis Agent: {e}")
#         return [], context

# # --- Function to use the comment ---
# def add_comments_to_docx(original_path, corrections):
#     try:
#         doc = docx.Document(original_path)
#         commented_texts = set()

#         for correction in corrections:
#             original_text = correction.get("original_text", "").strip()
#             if not original_text or original_text in commented_texts:
#                 continue

#             for paragraph in doc.paragraphs:
#                 if original_text in paragraph.text and original_text not in commented_texts:
#                     comment_text = (
#                         f"Suggestion: {correction.get('corrected_text', 'N/A')}\n\n"
#                         f"Issue(s) Found: {', '.join(correction.get('issues_found', []))}\n\n"
#                         f"Reason: {correction.get('consolidated_reason', 'N/A')}\n"
#                     )
                    
#                     doc.add_comment(
#                         text=comment_text,
#                         runs=paragraph.runs,
#                         author="AI Reviewer",
#                         initials="AI"
#                     )
                    
#                     commented_texts.add(original_text)
#                     break 
        
#         original_filename = os.path.basename(original_path)
#         name, ext = os.path.splitext(original_filename)
#         new_filename = f"{name}_reviewed{ext}"
#         save_path = os.path.join(GENERATED_DIR, new_filename)
        
#         doc.save(save_path)
#         print(f"Successfully saved commented document to {save_path}")
#         return new_filename
#     except Exception as e:
#         print(f"Error adding comments to DOCX: {e}")
#         return None
    
# # --- FLASK ROUTES ---
# @app.route('/')
# def index():
#     return render_template("ui_template.html")

# @app.route('/review_document', methods=['POST'])
# def review_document():
#     if llm is None: return jsonify({"error": "AI models not loaded."}), 503
#     if 'document_file' not in request.files: return jsonify({"error": "No document file part"}), 400
#     file = request.files['document_file']
#     if file.filename == '' or not allowed_file(file.filename): return jsonify({"error": "Invalid or no file selected"}), 400

#     filename = secure_filename(file.filename)
#     file_path = os.path.join(UPLOADS_DIR, filename)
#     file.save(file_path)

#     is_docx = filename.lower().endswith(".docx")
#     raw_text = ""

#     if filename.lower().endswith(".pdf"): raw_text = extract_text_from_pdf_fitz(file_path)
#     elif is_docx: raw_text = extract_text_from_docx(file_path)
#     else:
#         with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#             raw_text = f.read()

#     if not raw_text.strip(): 
#         os.remove(file_path) # Clean up if text extraction fails
#         return jsonify({"error": "Could not extract text."}), 400

#     # Batching Logic
#     paragraphs = [p.strip() for p in raw_text.split('\n\n') if len(p.strip()) > 20]
    
#     MAX_CHARS_PER_BATCH = 1000 
    
#     paragraph_batches = []
#     current_batch = []
#     current_chars = 0

#     for paragraph in paragraphs:
#         # If a single paragraph is too large, process it alone.
#         if len(paragraph) > MAX_CHARS_PER_BATCH:
#             # If a batch is already in progress, finalize it.
#             if current_batch:
#                 paragraph_batches.append("\n\n".join(current_batch))
#                 current_batch = []
#                 current_chars = 0
#             # Add the huge paragraph as its own batch.
#             paragraph_batches.append(paragraph)
#             continue

#         # If adding the next paragraph fits, add it to the current batch.
#         if current_chars + len(paragraph) <= MAX_CHARS_PER_BATCH:
#             current_batch.append(paragraph)
#             current_chars += len(paragraph)
#         # Otherwise, finalize the current batch and start a new one.
#         else:
#             paragraph_batches.append("\n\n".join(current_batch))
#             current_batch = [paragraph]
#             current_chars = len(paragraph)

#     # Add the last remaining batch if it exists.
#     if current_batch:
#         paragraph_batches.append("\n\n".join(current_batch))
    
#     all_corrections = []
#     for batch_text in paragraph_batches:
#         # Step 1: Generate Candidate Corrections (Your existing function)
#         candidate_corrections, relevant_rules_text = generate_review_for_batch(batch_text) # Modified to return rules
        
#         if not candidate_corrections:
#             continue

#         # Step 2: Verification Loop
#         all_corrections.extend(candidate_corrections)
        
#     download_filename = None
#     if is_docx and all_corrections:
#         download_filename = add_comments_to_docx(file_path, all_corrections)
#     os.remove(file_path)
    
#     # Prepare the response
#     response_data = {"corrections": all_corrections}
#     if download_filename:
#         response_data["download_url"] = f"/download/{download_filename}"
    
#     return jsonify(response_data)

# # --- Route to handle file downloads ---
# @app.route('/download/<filename>')
# def download_file(filename):
#     return send_from_directory(GENERATED_DIR, filename, as_attachment=True)

# if __name__ == "__main__":
#     app.run(host="0.0.0.0", port=5002, debug=True)

###########################################################################################
#############################################################################################


import os
import json
import re
import base64
from flask import Flask, request, jsonify, send_from_directory, render_template
import docx
from docx.document import Document
from docx.oxml.ns import qn
import requests
from werkzeug.utils import secure_filename

# --- Global App Settings ---
ollama_model_name = r"gemma3:12b" 
OLLAMA_API_URL = "http://localhost:11434/api/generate"

app = Flask(__name__)
UPLOADS_DIR = "uploads"
GENERATED_DIR = "generated"
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

# --- 1. OLLAMA CLIENT (Vision Enabled) ---
def call_ollama_vision(prompt, images=None):
    """
    Sends Text + Images to Ollama.
    """
    payload = {
        "model": ollama_model_name,
        "prompt": prompt,
        "stream": False,
        "options": {
            "num_ctx": 8192,
            "temperature": 0.1, # Low temperature for strict editing
            "num_predict": 2048,
        }
    }
    if images:
        payload["images"] = images

    try:
        response = requests.post(OLLAMA_API_URL, json=payload)
        response.raise_for_status()
        return response.json().get('response', '')
    except Exception as e:
        print(f"Error calling Ollama API: {e}")
        return "[]"

# --- 2. DOCX & IMAGE EXTRACTION HELPER ---
# def get_image_from_run(run, parent_part):
#     """Extracts base64 image from a run if present."""
#     try:
#         blip_matches = run._element.xpath('.//a:blip')
#         if not blip_matches: return None
#         for blip in blip_matches:
#             embed_attr = blip.get(qn('r:embed'))
#             if not embed_attr: continue
#             image_part = parent_part.related_parts.get(embed_attr)
#             if image_part:
#                 return base64.b64encode(image_part.blob).decode('utf-8')
#     except: pass
#     return None

def get_image_from_run(run, parent_part):
    """Extracts base64 image from a run if present."""
    try:
        # FIX: Use local-name() to ignore namespace prefixes (e.g. 'a:blip')
        # This avoids the "Undefined namespace prefix" error since we can't pass a namespace dict.
        blip_matches = run._element.xpath('.//*[local-name()="blip"]')
        if not blip_matches: return None
        
        for blip in blip_matches:
            # Use qn to get the fully qualified name for r:embed
            embed_attr = blip.get(qn('r:embed'))
            if not embed_attr: continue
            
            image_part = parent_part.related_parts.get(embed_attr)
            if image_part:
                return base64.b64encode(image_part.blob).decode('utf-8')
    except Exception as e:
        print(f"Image extraction warning: {e}")
    return None

def get_formatting_tags(paragraph):
    """
    Extracts indentation and numbering info to help AI detect formatting errors.
    """
    tags = []
    
    # 1. Indentation Check
    if paragraph.paragraph_format.left_indent:
        indent_val = paragraph.paragraph_format.left_indent.inches
        # Ignore small/standard indentations (like 0.5 for first line)
        if indent_val and indent_val > 0.6: 
            tags.append(f"[INDENTATION: {indent_val:.1f} inches]")

    # 2. Manual Numbering Check (Regex)
    # Detects start of line patterns like "1.", "1.2", "A)", "(i)"
    text = paragraph.text.strip()
    if re.match(r'^(\d+(\.\d+)*|[A-Z]|[a-z]|\([a-z0-9]+\))[\.\)]\s', text):
        tags.append("[NUMBERED_ITEM]")
    
    # 3. Automatic List Style Check (Word XML)
    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        tags.append("[AUTO_LIST_ITEM]")

    return " ".join(tags)

def table_to_markdown(table):
    """Converts table to Markdown to expose empty cells."""
    md_rows = []
    try:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_rows.append("| " + " | ".join(headers) + " |")
        md_rows.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            # Explicitly mark empty cells so LLM can catch them
            row_cells = [cell.text.strip() if cell.text.strip() else "MISSING_DATA" for cell in row.cells]
            md_rows.append("| " + " | ".join(row_cells) + " |")
    except: return "[Complex Table]"
    return "\n".join(md_rows)

def iter_block_items(parent):
    if isinstance(parent, Document): parent_elm = parent.element.body
    else: parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'): yield docx.text.paragraph.Paragraph(child, parent)
        elif child.tag.endswith('tbl'): yield docx.table.Table(child, parent)

def analyze_doc_structure(file_path):
    doc = docx.Document(file_path)
    blocks = []
    
    blank_line_counter = 0

    for block in iter_block_items(doc):
        # A. Paragraphs
        if isinstance(block, docx.text.paragraph.Paragraph):
            text = block.text.strip()
            style = block.style.name.lower()
            
            # Feature 5: Check for Extra Spaces (Python Logic)
            if not text:
                blank_line_counter += 1
                if blank_line_counter > 2:
                    blocks.append({
                        "type": "issue",
                        "original_text": "[BLANK SPACE]",
                        "corrected_text": "[DELETE]",
                        "issues_found": ["Formatting"],
                        "consolidated_reason": "Excessive vertical whitespace (3+ empty lines)."
                    })
                continue 
            else:
                blank_line_counter = 0

            # 1. Formatting Detection
            fmt_tags = get_formatting_tags(block)

            # Image Extraction
            images = []
            for run in block.runs:
                img = get_image_from_run(run, doc.part)
                if img: images.append(img)
            
            blocks.append({
                "type": "paragraph",
                "text": text,
                "style": style,
                "images": images
            })

        # B. Tables
        elif isinstance(block, docx.table.Table):
            blocks.append({
                "type": "table",
                "text": table_to_markdown(block),
                "images": []
            })
            blank_line_counter = 0

    return blocks

# --- 3. CORE ANALYSIS LOGIC (Prompt Engineering) ---
def run_gemma_analysis(content_blocks):
    all_corrections = []
    
    BATCH_SIZE = 5
    for i in range(0, len(content_blocks), BATCH_SIZE):
        batch = content_blocks[i:i+BATCH_SIZE]
        
        # 1. Add Python-detected issues immediately
        python_found_issues = [b for b in batch if b.get("type") == "issue"]
        all_corrections.extend(python_found_issues)
        
        llm_batch = [b for b in batch if b.get("type") != "issue"]
        if not llm_batch: continue

        prompt_text = ""
        batch_images = []
        img_idx = 1
        
        for item in llm_batch:
            if item['images']:
                prompt_text += f"\n[IMAGE {img_idx}]\n"
                prompt_text += f"Surrounding Text: {item['text']}\n"
                batch_images.extend(item['images'])
                img_idx += 1
            elif item['type'] == 'table':
                prompt_text += f"\n[TABLE START]\n{item['text']}\n[TABLE END]\n"
            else:
                prefix = f"[{item['style'].upper()}]" 
                prompt_text += f"{prefix}: {item['text']}\n"

        if not prompt_text.strip(): continue

        print(f"Processing batch with {len(batch_images)} images...")

        system_prompt = """You are a professional Technical Editor. Review the provided document segment.

**STRICT EDITORIAL RULES:**
1. **Grammar & Flow:** Correct grammar, punctuation, and awkward phrasing.
2. **Missing Captions:** If you see an [IMAGE X] or [TABLE], check if there is a caption immediately before or after it.
3. **Formatting:** Ensure Headings are logical.
4. **Table Data:** If a cell contains 'MISSING_DATA' or '||', flag it.
5. **Formatting:** - Check [INDENTATION]: Is it random or inconsistent with the text style?
    - Check [NUMBERED_ITEM]: Is the sequence logical? (e.g. "1." followed by "3." or "A." followed by "C.").

**OUTPUT FORMAT:**
Return ONLY a valid JSON list. Escape all backslashes (e.g. use \\\\ for \\).
[
  {
    "original_text": "text",
    "corrected_text": "correction",
    "issues_found": ["Grammar"],
    "consolidated_reason": "reason"
  }
]
If no errors, return [].
"""
        full_prompt = f"<|start_header_id|>system<|end_header_id|>\n\n{system_prompt}<|eot_id|><|start_header_id|>user<|end_header_id|>\n\n{prompt_text}<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n"
        
        response_text = call_ollama_vision(full_prompt, batch_images)
        
        # --- IMPROVED JSON PARSING ---
        try:
            # 1. Basic Clean
            clean_json = response_text.replace("```json", "").replace("```", "").strip()
            
            # 2. Try parsing
            corrections = json.loads(clean_json)
            all_corrections.extend(corrections)
            
        except json.JSONDecodeError:
            # 3. First Fallback: Fix Control Characters (newlines in strings)
            try:
                # Replace actual newlines with \n inside the string
                clean_json_fixed = clean_json.replace('\n', '\\n')
                corrections = json.loads(clean_json_fixed)
                all_corrections.extend(corrections)
                print(" -> Successfully repaired JSON (Control Characters)")
            except:
                # 4. Second Fallback: Fix Backslashes (Windows paths)
                try:
                    # Escape backslashes that aren't already escaped
                    clean_json_fixed = re.sub(r'(?<!\\)\\(?!["\\/bfnrtu])', r'\\\\', clean_json)
                    corrections = json.loads(clean_json_fixed)
                    all_corrections.extend(corrections)
                    print(" -> Successfully repaired JSON (Backslashes)")
                except:
                    print(f" -> Skipped batch due to malformed JSON from AI.")
                    # Optional: print(clean_json) to see what went wrong
        except Exception as e:
            print(f"General Error: {e}")

    return all_corrections

# --- 4. COMMENT INJECTION (Logic from app.py) ---
# def add_comments_to_docx(original_path, corrections):
#     try:
#         doc = docx.Document(original_path)
#         commented_texts = set()

#         for correction in corrections:
#             original_text = correction.get("original_text", "").strip()
#             if not original_text or original_text in commented_texts:
#                 continue

#             for paragraph in doc.paragraphs:
#                 if original_text in paragraph.text and original_text not in commented_texts:
#                     comment_text = (
#                         f"Suggestion: {correction.get('corrected_text', 'N/A')}\n\n"
#                         f"Issue(s) Found: {', '.join(correction.get('issues_found', []))}\n\n"
#                         f"Reason: {correction.get('consolidated_reason', 'N/A')}\n"
#                     )
                    
#                     # Assuming custom python-docx environment as per user request
#                     doc.add_comment(
#                         text=comment_text,
#                         runs=paragraph.runs,
#                         author="AI Reviewer",
#                         initials="AI"
#                     )
                    
#                     commented_texts.add(original_text)
#                     break 
        
#         original_filename = os.path.basename(original_path)
#         name, ext = os.path.splitext(original_filename)
#         new_filename = f"{name}_reviewed{ext}"
#         save_path = os.path.join(GENERATED_DIR, new_filename)
        
#         doc.save(save_path)
#         print(f"Successfully saved commented document to {save_path}")
#         return new_filename
#     except Exception as e:
#         print(f"Error adding comments to DOCX: {e}")
#         return None

# --- 4. IMPROVED COMMENT INJECTION ---
def add_comments_to_docx(original_path, corrections):
    try:
        doc = docx.Document(original_path)
        commented_texts = set()

        for correction in corrections:
            original_text = correction.get("original_text", "").strip()
            
            # --- CASE A: BLANK SPACES (Python Issue) ---
            if original_text == "[BLANK SPACE]":
                empty_count = 0
                for paragraph in doc.paragraphs:
                    if not paragraph.text.strip():
                        empty_count += 1
                        # If we hit 3 empty lines, comment on the 3rd one
                        if empty_count == 3: 
                            doc.add_comment(
                                text=f"Suggestion: {correction.get('corrected_text')}\nReason: {correction.get('consolidated_reason')}",
                                runs=paragraph.runs if paragraph.runs else paragraph.add_run(" "), # Ensure there is a run to attach to
                                author="AI Reviewer",
                                initials="AI"
                            )
                            # Only comment on the first instance of a block to avoid spam
                            # If you want to comment on ALL instances, remove the break logic below and improve 'commented_texts' handling
                            break 
                    else:
                        empty_count = 0
                continue

            # --- CASE B: TABLE DATA ---
            # If the issue mentions "Missing Data" or the text looks like a Markdown table row
            is_table_issue = "MISSING_DATA" in original_text or "|" in original_text
            
            if is_table_issue:
                # 1. Clean the markdown to get keywords (e.g. "| 102 | Bob |" -> ["102", "Bob"])
                keywords = [k.strip() for k in original_text.split('|') if k.strip() and k.strip() != "MISSING_DATA"]
                
                # 2. Search all tables
                for table in doc.tables:
                    for row in table.rows:
                        # Convert row to string to check if it matches keywords
                        row_text = " ".join([cell.text for cell in row.cells])
                        
                        # Check if all keywords are in this row (Fuzzy Match)
                        if keywords and all(k in row_text for k in keywords):
                            if row_text in commented_texts: continue
                            
                            # Find the empty cell to attach the comment
                            target_paragraph = row.cells[0].paragraphs[0] # Default to first cell
                            for cell in row.cells:
                                if not cell.text.strip():
                                    if cell.paragraphs:
                                        target_paragraph = cell.paragraphs[0]
                                        break
                            
                            doc.add_comment(
                                text=f"Suggestion: {correction.get('corrected_text')}\nReason: {correction.get('consolidated_reason')}",
                                runs=target_paragraph.runs if target_paragraph.runs else target_paragraph.add_run(" "),
                                author="AI Reviewer",
                                initials="AI"
                            )
                            commented_texts.add(row_text)
                            break 
                continue

            # --- CASE C: STANDARD TEXT (Body + Captions) ---
            if original_text in commented_texts: continue

            # 1. Search Body Paragraphs
            found = False
            for paragraph in doc.paragraphs:
                if original_text in paragraph.text:
                    doc.add_comment(
                        text=f"Suggestion: {correction.get('corrected_text')}\nReason: {correction.get('consolidated_reason')}",
                        runs=paragraph.runs,
                        author="AI Reviewer",
                        initials="AI"
                    )
                    commented_texts.add(original_text)
                    found = True
                    break 
            
            if found: continue

            # 2. Search Table Paragraphs (Text inside tables)
            # (Standard loop misses text inside tables, so we must loop explicitly)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if original_text in paragraph.text:
                                doc.add_comment(
                                    text=f"Suggestion: {correction.get('corrected_text')}\nReason: {correction.get('consolidated_reason')}",
                                    runs=paragraph.runs,
                                    author="AI Reviewer",
                                    initials="AI"
                                )
                                commented_texts.add(original_text)
                                found = True
                                break
                        if found: break
                    if found: break
        
        new_filename = f"Reviewed_{os.path.basename(original_path)}"
        save_path = os.path.join(GENERATED_DIR, new_filename)
        doc.save(save_path)
        return new_filename

    except Exception as e:
        print(f"Error adding comments to DOCX: {e}")
        return None



# --- 5. ROUTES ---
@app.route('/')
def index():
    return render_template("ui_template.html")

@app.route('/review_document', methods=['POST'])
def review_document():
    # 1. Check if file part exists
    if 'document_file' not in request.files: 
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['document_file']
    
    # 2. Check filename and extension
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
        
    filename = secure_filename(file.filename)
    
    # --- CRITICAL FIX: Block PDFs and other formats ---
    if not filename.lower().endswith('.docx'):
        return jsonify({
            "error": "Format not supported. Please upload a .docx file. (PDFs do not contain the structural data needed for this specific analysis)"
        }), 400
    # --------------------------------------------------

    file_path = os.path.join(UPLOADS_DIR, filename)
    file.save(file_path)

    try:
        # 3. Analyze Structure & Images
        content_blocks = analyze_doc_structure(file_path)
        
        # 4. Run Gemma 3
        corrections = run_gemma_analysis(content_blocks)
        
        # 5. Inject Comments (Using add_comments_to_docx logic)
        download_filename = add_comments_to_docx(file_path, corrections)
        
        return jsonify({
            "corrections": corrections,
            "download_url": f"/download/{download_filename}"
        })
        
    except Exception as e:
        print(f"Processing Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(GENERATED_DIR, filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5002)